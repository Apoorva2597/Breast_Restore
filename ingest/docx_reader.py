from __future__ import annotations
from pathlib import Path
from typing import List
from docx import Document
from ..models import NoteDocument

def _table_to_text(table) -> str:
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            t = (cell.text or "").strip()
            if t:
                cells.append(t.replace("\n", " ").strip())
        if cells:
            lines.append(" | ".join(cells))
    return "\n".join(lines)

def read_docx(path: str, note_id: str | None = None) -> NoteDocument:
    p = Path(path)
    doc = Document(str(p))
    parts: List[str] = []

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)

    for table in doc.tables:
        tt = _table_to_text(table)
        if tt.strip():
            parts.append("\n[TABLE]\n" + tt)

    text = "\n".join(parts).strip()
    return NoteDocument(note_id=note_id or p.stem, text=text, metadata={"source_path": str(p)})
